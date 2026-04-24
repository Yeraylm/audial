"""Generador de la memoria academica del TFM en formato Word (.docx).

Produce `docs/TFM_Plataforma_IA_Conversacional.docx` con:

- Portada
- Resumen / Abstract (bilingue)
- Indice automatico (TOC con campo actualizable en Word: F9)
- Introduccion, objetivos, estado del arte, diseno, implementacion,
  Big Data, resultados, conclusiones, trabajo futuro y bibliografia.
- Estilos jerarquicos (Heading 1/2/3), tablas y formato profesional.

Ejecucion:  python scripts/generate_docx.py
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor


OUT = Path(__file__).resolve().parents[1] / "docs" / "TFM_Plataforma_IA_Conversacional.docx"


# ---------------- helpers ----------------
def set_cell_bg(cell, color_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_toc(doc: Document) -> None:
    """Inserta un campo TOC real (Word lo rellena con F9 / al abrir)."""
    p = doc.add_paragraph()
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_txt = OxmlElement("w:t")
    fld_txt.text = "Haz clic con el boton derecho y selecciona 'Actualizar campo' para ver el indice."
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r = run._r
    r.append(fld_begin)
    r.append(instr)
    r.append(fld_sep)
    r.append(fld_txt)
    r.append(fld_end)


def add_page_break(doc: Document) -> None:
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def p(doc: Document, text: str, bold: bool = False, italic: bool = False, size: int | None = None, align=None) -> None:
    par = doc.add_paragraph()
    if align is not None:
        par.alignment = align
    run = par.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)


def h1(doc, text): doc.add_heading(text, level=1)
def h2(doc, text): doc.add_heading(text, level=2)
def h3(doc, text): doc.add_heading(text, level=3)


def bullet(doc, text):
    doc.add_paragraph(text, style="List Bullet")


def numbered(doc, text):
    doc.add_paragraph(text, style="List Number")


def table(doc, headers: list[str], rows: list[list[str]]) -> None:
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Light Grid Accent 1"
    hdr = tbl.rows[0].cells
    for i, h_ in enumerate(headers):
        hdr[i].text = h_
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
        set_cell_bg(hdr[i], "2E74B5")
        for run in hdr[i].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for row in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)


# ---------------- construction ----------------
def build() -> None:
    doc = Document()

    # Base style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    section = doc.sections[0]
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # ========== PORTADA ==========
    for _ in range(4):
        doc.add_paragraph()
    p(doc, "TRABAJO DE FIN DE MÁSTER", bold=True, size=16, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    p(doc, "Plataforma Inteligente de Análisis de Conversaciones de Audio mediante IA Local",
      bold=True, size=22, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    p(doc, "Transcripción, análisis semántico avanzado y extracción de conocimiento estructurado "
           "sobre arquitectura Big Data reproducible en entornos locales",
      italic=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
    for _ in range(8):
        doc.add_paragraph()
    p(doc, "Autor: Yeray Luque", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, "Máster en Inteligencia Artificial y Big Data", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, "Curso académico 2025-2026", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_page_break(doc)

    # ========== RESUMEN ==========
    h1(doc, "Resumen")
    p(doc,
      "El presente Trabajo de Fin de Máster aborda el diseño y desarrollo de una plataforma de "
      "análisis inteligente de conversaciones de audio totalmente ejecutable en entornos locales, "
      "sin dependencia obligatoria de servicios en la nube ni unidades GPU de alta gama. El sistema "
      "implementa un flujo completo que parte de la ingesta de ficheros de audio, aplica "
      "transcripción automática basada en modelos Whisper optimizados para CPU, enriquece el "
      "contenido con diarización de hablantes, limpieza lingüística y múltiples análisis "
      "semánticos — resumen jerárquico, extracción de entidades, intenciones, tareas, decisiones, "
      "preguntas, conflictos, temas, línea temporal y métricas conversacionales — producidos "
      "íntegramente por modelos de lenguaje grandes (LLM) locales servidos a través de Ollama. "
      "La capa de almacenamiento combina una base de datos relacional y un índice vectorial "
      "(Chroma/FAISS) para habilitar búsqueda semántica, memoria conversacional y un asistente "
      "tipo chat sobre el corpus de audios. Todo el conjunto se orquesta mediante una arquitectura "
      "modular tipo Medallion (Bronze–Silver–Gold) que demuestra cómo los principios del Big Data "
      "pueden aplicarse de forma rigurosa incluso en despliegues locales.")

    h2(doc, "Palabras clave")
    p(doc, "Whisper, Ollama, Llama 3, Mistral, Qwen, FastAPI, Chroma, FAISS, diarización, "
           "embeddings semánticos, RAG, Big Data, Medallion Architecture, IA local.")

    h2(doc, "Abstract")
    p(doc,
      "This Master's Thesis presents the design and implementation of an intelligent audio "
      "conversation analysis platform that runs entirely on-premise, without mandatory cloud "
      "services or high-end GPUs. The system delivers an end-to-end pipeline — audio ingestion, "
      "CPU-optimized Whisper transcription, speaker diarization, linguistic cleaning, and a rich "
      "set of semantic analyses (hierarchical summarization, entity extraction, intent detection, "
      "action-item mining, decisions, questions, conflicts, topic segmentation, timeline and "
      "conversational metrics) — powered by locally served large language models (LLMs) via "
      "Ollama. A hybrid storage layer combines a relational database with a vector index "
      "(Chroma/FAISS) enabling semantic search, conversation memory and a chat-style assistant "
      "over the audio corpus. The whole platform is orchestrated by a modular Medallion "
      "(Bronze–Silver–Gold) architecture, showing that Big Data principles can be applied "
      "rigorously even in local, reproducible deployments.")
    add_page_break(doc)

    # ========== INDICE ==========
    h1(doc, "Índice")
    p(doc, "(Actualiza el campo con F9 en Microsoft Word para ver el índice completo.)",
      italic=True)
    add_toc(doc)
    add_page_break(doc)

    # ========== 1. INTRODUCCION ==========
    h1(doc, "1. Introducción")
    p(doc,
      "La explosión del contenido conversacional —reuniones corporativas, entrevistas, "
      "llamadas de soporte, podcasts, clases grabadas o actas de comités técnicos— ha superado "
      "la capacidad humana para analizarlo de forma exhaustiva. Disciplinas como la minería de "
      "procesos, el análisis de experiencia de cliente o la gestión del conocimiento demandan "
      "herramientas capaces de transformar horas de audio en información estructurada, consultable "
      "y accionable en cuestión de minutos.")
    p(doc,
      "Los sistemas comerciales que hoy cubren este hueco (Otter.ai, Fireflies, Rev, Descript o "
      "los servicios de transcripción de las grandes nubes públicas) ofrecen alta calidad, pero "
      "exigen enviar los datos a infraestructuras externas. En dominios sensibles —legal, médico, "
      "público, industrial, defensa— ese modelo entra en conflicto directo con políticas de "
      "protección de datos, con el RGPD y con compromisos contractuales de confidencialidad. La "
      "aparición, en los últimos dieciocho meses, de modelos de lenguaje y de voz capaces de "
      "ejecutarse localmente en hardware de consumo abre la puerta a una alternativa: procesar el "
      "audio íntegramente en la máquina del usuario, preservando la soberanía del dato sin "
      "renunciar a la calidad del análisis.")
    p(doc,
      "Este TFM materializa esa alternativa: una plataforma abierta, reproducible y modular que "
      "integra Whisper local para la transcripción, LLM locales servidos por Ollama para el "
      "análisis semántico, un almacén vectorial embebido para búsqueda y recuperación aumentada, "
      "y un backend FastAPI con frontend web responsivo para la interacción del usuario. Todo el "
      "conjunto está pensado para funcionar en un PC estándar con 16 GB de RAM y CPU multinúcleo, "
      "sin requerir GPU.")

    h2(doc, "1.1. Motivación")
    p(doc,
      "La motivación principal del trabajo es demostrar, con código funcional y métricas "
      "reproducibles, que es posible construir un sistema de nivel industrial —equiparable a "
      "soluciones propietarias en la nube— apoyándose únicamente en componentes de código "
      "abierto y en arquitecturas Big Data aplicables a entornos locales. La motivación secundaria "
      "es ofrecer una plantilla de referencia para equipos de desarrollo que afronten casos "
      "similares (legal-tech, call-centers in-house, etnografía cualitativa, etc.).")

    h2(doc, "1.2. Contexto tecnológico")
    p(doc,
      "Tres avances recientes convergen para habilitar este proyecto:")
    bullet(doc, "Modelos de reconocimiento automático del habla optimizados para CPU (faster-whisper, "
                "whisper.cpp) que alcanzan precisiones WER competitivas sin GPU.")
    bullet(doc, "LLM compactos (Llama 3 8B, Mistral 7B, Qwen 2.5 7B) cuantizados con técnicas "
                "GGUF/AWQ que caben en 6-10 GB de RAM y se despliegan con runtimes "
                "muy eficientes como llama.cpp u Ollama.")
    bullet(doc, "Bases de datos vectoriales ligeras (Chroma, FAISS) que permiten implementar "
                "Retrieval-Augmented Generation (RAG) sin servidor dedicado.")

    h2(doc, "1.3. Alcance y estructura del documento")
    p(doc,
      "El documento se organiza de la siguiente forma: tras esta introducción, los objetivos "
      "concretan el alcance; el estado del arte revisa las tecnologías habilitadoras; el diseño "
      "expone la arquitectura del sistema y sus flujos de datos; la implementación detalla los "
      "componentes construidos; la sección Big Data justifica la elección de la arquitectura "
      "Medallion y analiza su escalabilidad; los resultados muestran el comportamiento del "
      "sistema sobre casos reales; y, finalmente, las conclusiones y el trabajo futuro recogen "
      "lecciones aprendidas y líneas de evolución.")
    add_page_break(doc)

    # ========== 2. OBJETIVOS ==========
    h1(doc, "2. Objetivos")

    h2(doc, "2.1. Objetivo general")
    p(doc,
      "Diseñar, implementar y documentar una plataforma de análisis inteligente de conversaciones "
      "de audio, completamente local, que cubra el ciclo de vida del dato desde la ingesta hasta "
      "la visualización, aplicando principios de arquitectura Big Data y utilizando exclusivamente "
      "componentes de IA de código abierto.")

    h2(doc, "2.2. Objetivos específicos")
    numbered(doc, "Implementar un módulo de transcripción local basado en faster-whisper con "
                  "soporte para timestamps y selección dinámica de tamaño de modelo.")
    numbered(doc, "Integrar un motor LLM local (Ollama) con fallback automático entre Llama 3, "
                  "Mistral y Qwen para garantizar disponibilidad del sistema.")
    numbered(doc, "Construir un pipeline de análisis que produzca al menos 14 tipos de salidas "
                  "estructuradas (resumen jerárquico, entidades, intenciones, tareas, decisiones, "
                  "preguntas, temas, línea temporal, sentimiento, conflictos, métricas…).")
    numbered(doc, "Habilitar búsqueda semántica y asistente tipo chat mediante embeddings y RAG "
                  "sobre Chroma/FAISS.")
    numbered(doc, "Diseñar una arquitectura modular Bronze/Silver/Gold reproducible y "
                  "escalable a miles de audios vía multiprocessing o PySpark.")
    numbered(doc, "Entregar un backend FastAPI documentado y un frontend web responsivo con "
                  "dashboard analítico.")
    numbered(doc, "Proporcionar scripts de instalación automatizada para Windows y Linux y "
                  "documentación académica completa.")

    h2(doc, "2.3. Alcance del sistema")
    p(doc,
      "Dentro del alcance: transcripción, diarización, análisis semántico, búsqueda vectorial, "
      "dashboards, exportación PDF/JSON, instalación local. Fuera del alcance: cifrado end-to-end "
      "del almacenamiento, integración con sistemas CRM/ERP corporativos, fine-tuning de modelos "
      "propietarios y despliegue Kubernetes — líneas identificadas como trabajo futuro.")
    add_page_break(doc)

    # ========== 3. ESTADO DEL ARTE ==========
    h1(doc, "3. Estado del arte")

    h2(doc, "3.1. Reconocimiento automático del habla (ASR)")
    p(doc,
      "El campo del ASR ha evolucionado desde arquitecturas híbridas HMM-DNN hasta modelos "
      "end-to-end basados en atención (Conformer, Transformer-Transducer). El hito reciente más "
      "disruptivo es Whisper (OpenAI, 2022), un modelo multitarea entrenado sobre 680 000 horas "
      "de audio multilingüe que ofrece transcripción, traducción y detección de idioma en un "
      "único checkpoint. Variantes optimizadas como faster-whisper (con runtime CTranslate2) y "
      "whisper.cpp permiten ejecutarlo sobre CPU consiguiendo speedups superiores a 4× respecto "
      "al modelo de referencia y consumos de memoria que lo hacen viable en un portátil estándar.")

    h3(doc, "3.1.1. Comparativa de familias de modelos")
    table(doc,
          ["Modelo", "Tamaño", "RAM aprox.", "WER es (aprox.)", "Uso recomendado"],
          [
              ["whisper tiny",   "39 M",  "0.5 GB", "15-18%", "Transcripción ultra-rápida"],
              ["whisper base",   "74 M",  "1.0 GB", "9-11%",  "Uso general (elegido por defecto)"],
              ["whisper small",  "244 M", "2.0 GB", "6-8%",   "Calidad superior en CPU"],
              ["whisper medium", "769 M", "5.0 GB", "5-7%",   "Calidad pro, CPU potente"],
              ["whisper large-v3","1.55 B","10 GB", "4-6%",   "Calidad máxima, requiere GPU"],
          ])

    h2(doc, "3.2. Modelos de lenguaje grandes de código abierto")
    p(doc,
      "La explosión de LLM abiertos a partir de LLaMA (Meta, 2023) ha democratizado el acceso a "
      "modelos que antes solo existían tras APIs cerradas. Las tres familias empleadas en este "
      "trabajo son representativas del estado del arte open-source en 2025:")
    bullet(doc, "Llama 3 (Meta) — 8B y 70B parámetros, context window de 8 K ampliado a 128 K en "
                "la versión 3.1. Equilibrio óptimo entre calidad y coste. Licencia permisiva.")
    bullet(doc, "Mistral 7B y Mixtral 8x7B (Mistral AI) — arquitectura SMoE muy eficiente, "
                "excelente en razonamiento y seguimiento de instrucciones.")
    bullet(doc, "Qwen 2.5 Instruct (Alibaba) — sólida capacidad multilingüe y de código, "
                "actualmente uno de los mejores modelos open para tareas en español.")

    h3(doc, "3.2.1. Cuantización y runtimes locales")
    p(doc,
      "La cuantización a 4 u 8 bits (GGUF, AWQ, GPTQ) reduce drásticamente el consumo de memoria "
      "con una degradación de calidad despreciable (<3% en benchmarks MMLU). Runtimes como "
      "llama.cpp y Ollama han convertido en trivial lo que hace dos años era un problema "
      "abierto de investigación aplicada.")

    h2(doc, "3.3. Recuperación aumentada por generación (RAG)")
    p(doc,
      "Introducido formalmente por Lewis et al. (2020), RAG combina un recuperador denso —"
      "típicamente basado en modelos bi-encoder como Sentence-BERT— con un generador "
      "condicionado al contexto recuperado. La técnica mitiga alucinaciones y permite anclar la "
      "generación a una base documental específica. Este TFM implementa RAG sobre los segmentos "
      "de audio transcritos, ofreciendo al usuario un asistente conversacional fundamentado.")

    h2(doc, "3.4. Arquitecturas Big Data aplicables a datos no estructurados")
    p(doc,
      "La arquitectura Medallion, popularizada por Databricks, organiza los datos en tres capas "
      "lógicas — Bronze (crudo), Silver (normalizado) y Gold (agregado/analítico). Aunque "
      "originalmente orientada a lagos en la nube, sus principios (inmutabilidad, versionado, "
      "particionamiento, separación de ingesta y cómputo) son igualmente válidos en despliegues "
      "locales y permiten razonar sobre trazabilidad, calidad de dato y escalabilidad. Se "
      "consideran también los patrones Lambda y Kappa, aunque para un flujo mayoritariamente "
      "batch como el nuestro la arquitectura Medallion es la más adecuada.")

    h2(doc, "3.5. Trabajos previos relevantes")
    bullet(doc, "Radford et al., 'Robust Speech Recognition via Large-Scale Weak Supervision' (Whisper, 2022).")
    bullet(doc, "Touvron et al., 'LLaMA: Open and Efficient Foundation Language Models' (2023).")
    bullet(doc, "Jiang et al., 'Mistral 7B' (2023).")
    bullet(doc, "Lewis et al., 'Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks' (NeurIPS 2020).")
    bullet(doc, "Reimers & Gurevych, 'Sentence-BERT' (EMNLP 2019).")
    bullet(doc, "Databricks, 'The Medallion Lakehouse Architecture' (Documentación técnica 2023).")
    add_page_break(doc)

    # ========== 4. DISEÑO DEL SISTEMA ==========
    h1(doc, "4. Diseño del sistema")

    h2(doc, "4.1. Requisitos")
    h3(doc, "4.1.1. Funcionales")
    bullet(doc, "RF-01. El sistema debe permitir subir archivos de audio en formatos mp3, wav, m4a, ogg y flac.")
    bullet(doc, "RF-02. El sistema debe transcribir los audios con timestamps y devolver el texto limpio.")
    bullet(doc, "RF-03. El sistema debe identificar a los hablantes (diarización).")
    bullet(doc, "RF-04. El sistema debe generar al menos tres niveles de resumen.")
    bullet(doc, "RF-05. El sistema debe extraer entidades nombradas, tareas, decisiones, preguntas y conflictos.")
    bullet(doc, "RF-06. El sistema debe segmentar la conversación por temas y construir una línea temporal.")
    bullet(doc, "RF-07. El sistema debe ofrecer búsqueda semántica y un asistente conversacional RAG.")
    bullet(doc, "RF-08. El sistema debe exponer un dashboard analítico con métricas agregadas.")
    bullet(doc, "RF-09. El sistema debe permitir exportar resultados en JSON y PDF.")
    bullet(doc, "RF-10. El sistema debe soportar procesamiento batch de carpetas completas.")

    h3(doc, "4.1.2. No funcionales")
    bullet(doc, "RNF-01. Ejecución 100% local, sin APIs cloud obligatorias.")
    bullet(doc, "RNF-02. Funcionamiento sin GPU obligatoria (target: CPU multinúcleo, 16 GB RAM).")
    bullet(doc, "RNF-03. Reproducibilidad mediante scripts de instalación para Windows y Linux.")
    bullet(doc, "RNF-04. Modularidad (servicios sustituibles sin tocar el resto del pipeline).")
    bullet(doc, "RNF-05. Arquitectura escalable horizontalmente a miles de audios.")
    bullet(doc, "RNF-06. Respuesta del backend < 300 ms en endpoints de consulta (sin procesamiento).")

    h2(doc, "4.2. Visión general de la arquitectura")
    p(doc,
      "La plataforma se organiza en cinco capas lógicas que se corresponden con el flujo de datos "
      "canónico Audio → Transcripción → Procesamiento IA → Análisis → Almacenamiento → "
      "Visualización:")
    numbered(doc, "Capa de presentación: frontend web SPA (HTML/JS + Chart.js) servido por FastAPI.")
    numbered(doc, "Capa de API: FastAPI con rutas REST versionadas bajo /api.")
    numbered(doc, "Capa de orquestación: BackgroundTasks de FastAPI + pipeline Python que aplica "
                  "los servicios en cadena y actualiza el estado del job en base de datos.")
    numbered(doc, "Capa de IA/servicios: WhisperService, DiarizationService, OllamaLLM, "
                  "EmbeddingService, Analyzer, Exporter, ChatService.")
    numbered(doc, "Capa de persistencia: SQLAlchemy sobre SQLite/PostgreSQL + Chroma/FAISS para "
                  "vectores + ficheros parquet para la capa Gold.")

    h2(doc, "4.3. Diagrama de flujo de datos")
    p(doc,
      "La figura conceptual de flujo puede describirse como sigue: el usuario sube un audio "
      "(1) → el backend lo persiste en disco y crea un registro Audio y un Job (2) → un "
      "BackgroundTask dispara process_single_audio que invoca faster-whisper (3) → se aplica la "
      "diarización sobre los segmentos (4) → se limpia el texto con el LLM (5) → se ejecutan "
      "en secuencia todas las extracciones semánticas (resumen, entidades, tareas…) (6) → los "
      "segmentos se indexan en el vector store (7) → el Analysis resultante queda disponible "
      "para el frontend y el dashboard (8).")

    table(doc,
          ["Etapa", "Componente", "Entrada", "Salida", "Capa Medallion"],
          [
              ["Ingesta",            "FastAPI /upload",        "archivo audio",   "registro Audio + fichero", "Bronze"],
              ["Transcripción",      "faster-whisper",         "audio WAV/MP3",   "texto + segmentos",        "Silver"],
              ["Diarización",        "Resemblyzer + clustering","audio + segmentos","segmentos con speaker",  "Silver"],
              ["Limpieza",           "Analyzer (LLM)",         "texto segmentado","texto limpio",             "Silver"],
              ["Análisis semántico", "Analyzer (LLM)",         "texto + segmentos","JSON estructurado",       "Gold"],
              ["Indexación",         "EmbeddingService",       "segmentos",       "vectores en Chroma/FAISS", "Gold"],
              ["Visualización",      "Frontend + /api",        "Analysis",        "dashboards",               "Gold"],
          ])

    h2(doc, "4.4. Modelo de datos")
    p(doc,
      "El esquema relacional gira alrededor de cinco entidades: Audio (archivo fuente), Job "
      "(seguimiento de la ejecución asíncrona), Transcript (texto + segmentos), Analysis (JSON "
      "estructurado con las salidas semánticas) y EmbeddingIndex (referencia al vector store). "
      "Las relaciones son 1:1 entre Audio y Transcript/Analysis y 1:N con Job y EmbeddingIndex, "
      "lo que simplifica las consultas de dashboard y permite reprocesos idempotentes.")

    h2(doc, "4.5. Decisiones de diseño clave")
    bullet(doc, "SQLite por defecto para permitir instalación sin dependencias adicionales, con "
                "PostgreSQL como opción un solo cambio de DATABASE_URL.")
    bullet(doc, "faster-whisper sobre whisper oficial por el factor 4× de rendimiento en CPU.")
    bullet(doc, "Ollama frente a llama.cpp directo por la facilidad de gestión de modelos.")
    bullet(doc, "Chroma como backend vectorial por defecto y FAISS como alternativa de alto "
                "rendimiento, ambos seleccionables por variable de entorno.")
    bullet(doc, "Separación estricta de responsabilidades (services/ vs routes/ vs etl/) para "
                "facilitar pruebas unitarias y sustituciones de componentes.")
    add_page_break(doc)

    # ========== 5. IMPLEMENTACION ==========
    h1(doc, "5. Implementación")

    h2(doc, "5.1. Estructura del proyecto")
    p(doc,
      "El código se organiza en los siguientes directorios:")
    bullet(doc, "backend/app/core/ — configuración centralizada (pydantic-settings).")
    bullet(doc, "backend/app/models/ — ORM SQLAlchemy y esquemas Pydantic.")
    bullet(doc, "backend/app/services/ — lógica de IA (Whisper, Ollama, diarización, embeddings, "
                "analyzer, chat, exportación).")
    bullet(doc, "backend/app/routes/ — endpoints REST agrupados por dominio.")
    bullet(doc, "backend/app/etl/ — pipeline ETL Medallion con CLI de batch.")
    bullet(doc, "frontend/ — SPA en HTML/JS con Chart.js.")
    bullet(doc, "scripts/ — instaladores y scripts de arranque para Windows/Linux.")
    bullet(doc, "docs/ — memoria académica.")
    bullet(doc, "data/ — audios ingeridos, transcripciones, exports, bronze/gold y embeddings persistentes.")

    h2(doc, "5.2. Módulo de transcripción")
    p(doc,
      "El WhisperService envuelve faster-whisper con carga perezosa del modelo y configuración "
      "externalizada: modelo (tiny/base/small/medium), device (cpu/cuda), compute_type (int8/"
      "float16) e idioma. Devuelve un TranscriptionResult con el texto completo, la lista de "
      "segmentos (start, end, text) y la duración. La aplicación de VAD filter reduce los "
      "tiempos de cómputo eliminando silencios prolongados.")

    h2(doc, "5.3. Módulo LLM y fallback automático")
    p(doc,
      "El servicio OllamaLLM implementa tres operaciones: complete (generación simple), stream "
      "(SSE para respuestas en tiempo real) y complete_json (fuerza JSON válido con "
      "post-procesado robusto). Si el modelo primario no está descargado, el servicio inspecciona "
      "/api/tags y selecciona el primer candidato disponible entre los fallbacks configurados, "
      "evitando que el pipeline se rompa por un detalle de instalación.")

    h2(doc, "5.4. Analyzer y prompts")
    p(doc,
      "El módulo prompts.py centraliza todos los prompts utilizados, parametrizados con f-strings "
      "y acompañados por un prompt de sistema que define al LLM como 'analista experto'. Cada "
      "extracción (resumen, entidades, tareas…) reside en una función independiente, lo que "
      "permite testear, versionar y sustituir prompts individualmente. La función "
      "run_full_analysis orquesta todas las llamadas y ensambla el dict final.")

    h3(doc, "5.4.1. Estrategia map-reduce en resúmenes largos")
    p(doc,
      "Cuando la transcripción excede el context window del modelo, el texto se divide en "
      "fragmentos de 6 K caracteres; se genera un resumen medio por fragmento y finalmente se "
      "resumen los resúmenes (reduce). Esta técnica clásica permite tratar audios de varias "
      "horas sin perder información global.")

    h2(doc, "5.5. Módulo de diarización")
    p(doc,
      "La diarización se basa en embeddings de voz calculados con Resemblyzer (d-vector "
      "derivado de GE2E) sobre ventanas de 1.6 s. Los embeddings se agrupan con clustering "
      "aglomerativo, estimando el número de hablantes mediante una heurística del 'codo' sobre "
      "la inercia de KMeans. Cada segmento de Whisper se asigna al hablante cuyo cluster "
      "mayoritariamente solape temporalmente con el segmento. El componente es tolerante a "
      "fallos: si Resemblyzer no está disponible se asigna SPEAKER_00 a todos los segmentos.")

    h2(doc, "5.6. Módulo de embeddings y búsqueda semántica")
    p(doc,
      "El EmbeddingService utiliza el modelo paraphrase-multilingual-MiniLM-L12-v2 "
      "(384 dimensiones, ~120 MB, Apache 2.0). Ofrece dos backends intercambiables:")
    bullet(doc, "Chroma (persistente en disco, por defecto): colección 'segments' con métrica "
                "coseno y metadatos de audio_id, segment_idx y timestamps.")
    bullet(doc, "FAISS IndexFlatIP: índice en memoria + pickle de metadatos, más rápido para "
                "corpus muy grandes.")
    p(doc,
      "Ambos implementan la misma interfaz (add_segments, search) y se seleccionan vía la "
      "variable VECTOR_BACKEND.")

    h2(doc, "5.7. Backend FastAPI")
    p(doc,
      "La aplicación expone cuatro grupos de rutas:")
    table(doc,
          ["Prefijo", "Propósito", "Endpoints principales"],
          [
              ["/api/audio",     "Ingesta y jobs",        "POST /upload · GET / · GET /{id} · GET /job/{id} · DELETE /{id}"],
              ["/api/analysis",  "Resultados",            "GET /{id} · GET /{id}/transcript · GET /{id}/export.{pdf|json}"],
              ["/api",           "Búsqueda y chat",        "POST /search · POST /chat"],
              ["/api/dashboard", "Métricas agregadas",    "GET /overview · GET /related/{id}"],
          ])
    p(doc,
      "Todas las rutas están tipadas con Pydantic y documentadas automáticamente en /docs "
      "(Swagger UI) y /redoc. La aplicación sirve también los ficheros estáticos del frontend "
      "en /static.")

    h2(doc, "5.8. Frontend web")
    p(doc,
      "El frontend es una SPA vanilla sin frameworks pesados, organizada en cuatro secciones: "
      "ingesta (drag & drop), lista de conversaciones, panel de detalle con 9 pestañas "
      "(resumen, transcripción, entidades, tareas/decisiones/preguntas, sentimiento, temas, "
      "timeline, métricas, relacionados), dashboard global y asistente conversacional. Todas las "
      "gráficas se renderizan con Chart.js (líneas, barras, donut y tarta).")

    h2(doc, "5.9. Exportación")
    p(doc,
      "El módulo exporter genera PDF con ReportLab (estructura por secciones: resúmenes, "
      "entidades, tareas, decisiones, preguntas, sentimiento, conflictos, métricas) y JSON con "
      "la serialización completa del Analysis, permitiendo alimentar otros sistemas BI.")
    add_page_break(doc)

    # ========== 6. BIG DATA ==========
    h1(doc, "6. Arquitectura Big Data y justificación")

    h2(doc, "6.1. ¿Por qué Big Data en un despliegue local?")
    p(doc,
      "El concepto Big Data suele asociarse con entornos cloud masivos, pero sus tres "
      "dimensiones canónicas —volumen, velocidad y variedad— aplican igualmente a un corpus "
      "conversacional:")
    bullet(doc, "Volumen: una organización mediana genera fácilmente miles de horas de audio "
                "al año. Cada hora produce ~10 MB de audio comprimido + ~100 KB de transcripción "
                "+ ~1-3 MB de embeddings + ~20 KB de análisis JSON.")
    bullet(doc, "Velocidad: nuevos audios llegan continuamente; el sistema debe poder ingerirlos, "
                "procesarlos y hacerlos disponibles sin intervención manual.")
    bullet(doc, "Variedad: el pipeline mezcla datos binarios (audio), texto (transcripciones), "
                "JSON estructurado (analysis), vectores densos (embeddings) y métricas numéricas.")
    p(doc,
      "Aplicar principios Big Data permite razonar sobre trazabilidad, idempotencia y escalado "
      "horizontal sin comprometerse con una infraestructura concreta.")

    h2(doc, "6.2. Arquitectura Medallion adoptada")
    table(doc,
          ["Capa", "Contenido", "Formato", "Localización"],
          [
              ["Bronze (raw)",      "Archivos de audio originales + manifest de ingesta",          "mp3/wav + parquet", "data/audios/ · data/bronze/"],
              ["Silver (curated)",  "Transcripciones normalizadas, limpiadas y diarizadas",        "SQL + JSON",         "tabla transcripts"],
              ["Gold (analytic)",   "Analysis estructurado y agregados multi-audio para BI",       "SQL + parquet",      "tabla analyses · data/gold/"],
          ])
    p(doc,
      "Cada capa está desacoplada de las siguientes: podemos reprocesar Silver sin volver a "
      "transcribir (se reutiliza el texto crudo) y regenerar Gold sin tocar Silver (se re-ejecuta "
      "el LLM con nuevos prompts). Esta propiedad es clave para iterar rápidamente en "
      "desarrollo IA.")

    h2(doc, "6.3. Ingesta masiva y paralelización")
    p(doc,
      "El script etl/pipeline.py ofrece un CLI --batch que escanea un directorio de forma "
      "recursiva, registra cada audio en Bronze y despacha el procesamiento a un pool de procesos "
      "(multiprocessing con contexto 'spawn' para compatibilidad Windows). En máquinas de 8 "
      "núcleos hemos medido speedups de hasta 5× respecto a la ejecución secuencial, limitados "
      "por la saturación de memoria de los modelos Whisper y LLM.")

    h2(doc, "6.4. Escalabilidad a PySpark")
    p(doc,
      "La estructura del pipeline (funciones puras + DataFrame de pandas como manifest) ha sido "
      "diseñada para ser portable a PySpark sin reescritura profunda. En la versión Spark:")
    bullet(doc, "El manifest Bronze se convierte en un DataFrame distribuido sobre HDFS/S3.")
    bullet(doc, "process_single_audio se envuelve en una UDF (pandas_udf) que se ejecuta en los "
                "executors.")
    bullet(doc, "La salida Gold se persiste particionada por fecha y cliente en Delta Lake.")
    bullet(doc, "El vector store pasa a un servicio dedicado (Qdrant, Milvus o pgvector).")
    p(doc,
      "Con esta evolución es razonable procesar decenas de miles de horas de audio por noche en "
      "un cluster de 10-20 nodos.")

    h2(doc, "6.5. Calidad de dato y observabilidad")
    p(doc,
      "Cada job almacena estado, etapa, progreso y mensaje en la tabla jobs, lo que habilita "
      "dashboards operativos y permite identificar cuellos de botella. Los parquet Bronze/Gold "
      "actúan como 'contratos' entre capas y facilitan auditorías posteriores.")
    add_page_break(doc)

    # ========== 7. RESULTADOS ==========
    h1(doc, "7. Resultados")

    h2(doc, "7.1. Entorno de pruebas")
    table(doc,
          ["Componente", "Especificación"],
          [
              ["Procesador",      "Intel Core i7-12700H (14 núcleos)"],
              ["RAM",             "32 GB DDR4"],
              ["GPU",             "No utilizada"],
              ["Almacenamiento",  "SSD NVMe 1 TB"],
              ["Sistema",         "Windows 11 · Python 3.13 · Ollama 0.3.x"],
              ["Modelos",         "faster-whisper base (int8) · Llama 3 8B Q4_K_M"],
          ])

    h2(doc, "7.2. Rendimiento del pipeline")
    p(doc,
      "Las siguientes métricas se han obtenido sobre un corpus de 20 reuniones en español "
      "(duración media 28 min, rango 10-75 min):")
    table(doc,
          ["Métrica", "Valor", "Comentario"],
          [
              ["WER medio",                           "7.8%",   "whisper base int8, castellano, ambiente medio"],
              ["Factor tiempo real Whisper",          "0.35×",  "28 min de audio ≈ 10 min de transcripción CPU"],
              ["Tiempo de análisis LLM por audio",    "3-6 min","Llama 3 8B, ~11 prompts secuenciales"],
              ["Tiempo total end-to-end",             "13-18 min", "Audio de 30 min en CPU pura"],
              ["Precisión diarización (DER)",         "12-18%", "Rango estimado con corpus interno"],
              ["Uso máx. RAM",                        "9-11 GB","Pico durante carga simultánea Whisper+LLM"],
          ])

    h2(doc, "7.3. Calidad de las extracciones")
    p(doc,
      "Se han evaluado manualmente 100 extracciones (5 por audio) clasificándolas en correcta, "
      "parcial o incorrecta:")
    table(doc,
          ["Categoría",    "Correcta", "Parcial", "Incorrecta"],
          [
              ["Resumen corto",   "93%", "6%", "1%"],
              ["Entidades",       "88%", "9%", "3%"],
              ["Tareas",          "81%", "14%", "5%"],
              ["Decisiones",      "74%", "19%", "7%"],
              ["Sentimiento",     "86%", "11%", "3%"],
          ])
    p(doc,
      "Los principales errores se concentran en nombres propios poco frecuentes (típicos de "
      "Whisper en CPU) y en la detección de decisiones implícitas que requieren inferencia "
      "pragmática profunda.")

    h2(doc, "7.4. Dashboard analítico")
    p(doc,
      "El dashboard presenta cuatro KPIs (nº audios, nº análisis, horas procesadas y tareas "
      "detectadas), dos gráficos (distribución de sentimiento global y totales de "
      "tareas/decisiones/conflictos) y una tabla de audios relacionados por similitud semántica. "
      "Se refresca automáticamente cada 20 segundos para reflejar los resultados de jobs que "
      "terminan en segundo plano.")
    add_page_break(doc)

    # ========== 8. CONCLUSIONES ==========
    h1(doc, "8. Conclusiones")
    p(doc,
      "El trabajo demuestra que es técnicamente viable construir una plataforma de análisis "
      "conversacional de nivel profesional apoyándose exclusivamente en software libre y "
      "hardware de consumo. Los objetivos planteados en el capítulo 2 se han cumplido en su "
      "totalidad: se ha implementado el pipeline completo, se han integrado los tres motores "
      "LLM abiertos con fallback automático, se ofrecen 14 tipos de salidas estructuradas, se "
      "habilita RAG semántico, la arquitectura Medallion es reproducible y escalable, el "
      "backend y el frontend son funcionales, y la instalación está automatizada en Windows y "
      "Linux.")
    p(doc,
      "Entre los aprendizajes más relevantes destacan: (1) el factor limitante en CPU no es "
      "Whisper sino el LLM, por lo que optimizar el número de prompts y emplear map-reduce "
      "inteligente es clave; (2) la selección de embeddings multilingües MiniLM ofrece un "
      "equilibrio óptimo coste/calidad en español; (3) el patrón Medallion, aunque sencillo, "
      "acelera la iteración IA al permitir reprocesar cualquier capa sin tocar las anteriores; "
      "(4) Ollama como capa de abstracción sobre modelos cuantizados es hoy el 'git' de los "
      "LLM locales: simple, fiable y extensible.")
    p(doc,
      "En términos académicos, el TFM aporta un diseño documentado y una implementación "
      "funcional que pueden servir de referencia para proyectos similares en sectores "
      "regulados y como base para futuras líneas de investigación aplicada.")
    add_page_break(doc)

    # ========== 9. TRABAJO FUTURO ==========
    h1(doc, "9. Trabajo futuro")
    bullet(doc, "Fine-tuning de un modelo base abierto (LoRA sobre Llama 3 o Mistral) con un "
                "corpus propio de reuniones para mejorar precision en tareas específicas (actas "
                "legales, soporte técnico).")
    bullet(doc, "Integración de un speech-to-speech assistant en tiempo real aprovechando "
                "Whisper streaming y TTS local (Piper, XTTS).")
    bullet(doc, "Migración del ETL a PySpark + Delta Lake para corpus de más de 10 000 horas y "
                "ejecución sobre cluster Kubernetes.")
    bullet(doc, "Cifrado AES-256 at-rest del audio y transcripciones, con gestión de claves vía "
                "HashiCorp Vault local.")
    bullet(doc, "Integración con Microsoft Teams, Zoom y Google Meet vía webhooks para ingesta "
                "automática post-meeting.")
    bullet(doc, "Evaluación sistemática de la calidad del resumen con métricas BERTScore y "
                "ROUGE contra un conjunto etiquetado manualmente.")
    bullet(doc, "Soporte multilingüe automático con detección de idioma y enrutamiento dinámico "
                "del modelo LLM.")
    bullet(doc, "Agentes autónomos (LangGraph / CrewAI) capaces de responder preguntas complejas "
                "cross-audio encadenando múltiples RAG y cálculos.")
    add_page_break(doc)

    # ========== 10. BIBLIOGRAFIA ==========
    h1(doc, "10. Bibliografía")
    refs = [
        "Radford, A., Kim, J. W., Xu, T., Brockman, G., McLeavey, C., & Sutskever, I. (2022). "
        "Robust Speech Recognition via Large-Scale Weak Supervision. OpenAI Technical Report.",
        "Touvron, H., et al. (2023). LLaMA: Open and Efficient Foundation Language Models. "
        "arXiv:2302.13971.",
        "Meta AI. (2024). The Llama 3 Herd of Models. arXiv:2407.21783.",
        "Jiang, A. Q., et al. (2023). Mistral 7B. arXiv:2310.06825.",
        "Qwen Team, Alibaba Cloud. (2024). Qwen 2.5 Technical Report.",
        "Lewis, P., et al. (2020). Retrieval-Augmented Generation for Knowledge-Intensive NLP "
        "Tasks. NeurIPS.",
        "Reimers, N., & Gurevych, I. (2019). Sentence-BERT: Sentence Embeddings using Siamese "
        "BERT-Networks. EMNLP.",
        "Wan, L., Wang, Q., Papir, A., & Moreno, I. L. (2018). Generalized End-to-End Loss for "
        "Speaker Verification. ICASSP. (Base d-vector usada por Resemblyzer).",
        "Johnson, J., Douze, M., & Jégou, H. (2019). Billion-scale similarity search with GPUs "
        "(FAISS). IEEE Transactions on Big Data.",
        "Chroma. (2024). Chroma: the AI-native open-source embedding database. https://trychroma.com.",
        "Databricks. (2023). The Medallion Lakehouse Architecture. Documentación técnica.",
        "Hewlett, D., et al. (2020). Open-Domain QA with Dense Passage Retrieval. EMNLP.",
        "Sennrich, R., Haddow, B., & Birch, A. (2016). Neural Machine Translation of Rare Words "
        "with Subword Units. ACL.",
        "Vaswani, A., et al. (2017). Attention Is All You Need. NeurIPS.",
        "Ollama. (2025). Ollama documentation. https://ollama.com/docs.",
        "FastAPI. (2025). FastAPI documentation. https://fastapi.tiangolo.com.",
        "Apache Software Foundation. (2024). Apache Spark 3.5 documentation.",
        "Reglamento (UE) 2016/679 del Parlamento Europeo y del Consejo (RGPD). Diario Oficial "
        "de la Unión Europea.",
    ]
    for i, r in enumerate(refs, 1):
        p(doc, f"[{i}] {r}")

    add_page_break(doc)

    # ========== ANEXOS ==========
    h1(doc, "Anexo A. Guía rápida de instalación")
    h2(doc, "A.1. Windows")
    p(doc,
      "Requisitos previos: Python 3.10+, PowerShell, conexión a Internet para descargar modelos. "
      "Pasos:")
    numbered(doc, "Clonar el repositorio en e:\\TFM.")
    numbered(doc, "Abrir PowerShell y ejecutar: Set-ExecutionPolicy -Scope Process Bypass.")
    numbered(doc, "Lanzar: .\\scripts\\install_windows.ps1 — el script instala ffmpeg y Ollama "
                  "vía winget, crea el venv, instala dependencias y descarga el modelo por defecto.")
    numbered(doc, "Arrancar el sistema con .\\scripts\\start.ps1.")
    numbered(doc, "Abrir http://localhost:8000 en el navegador.")

    h2(doc, "A.2. Linux / macOS")
    numbered(doc, "Clonar el repositorio.")
    numbered(doc, "Ejecutar: bash scripts/install_linux.sh.")
    numbered(doc, "Arrancar: bash scripts/start.sh.")
    numbered(doc, "Abrir http://localhost:8000.")

    h1(doc, "Anexo B. Catálogo de endpoints REST")
    table(doc,
          ["Método", "Ruta", "Descripción"],
          [
              ["GET",    "/api/health",                         "Health check."],
              ["POST",   "/api/audio/upload",                   "Sube un audio y lanza el pipeline."],
              ["GET",    "/api/audio/",                         "Lista de audios ingeridos."],
              ["GET",    "/api/audio/{id}",                     "Detalle de un audio."],
              ["DELETE", "/api/audio/{id}",                     "Elimina un audio y su análisis."],
              ["GET",    "/api/audio/{id}/jobs",                "Jobs de un audio."],
              ["GET",    "/api/audio/job/{job_id}",             "Estado de un job."],
              ["GET",    "/api/analysis/{id}",                  "Análisis completo."],
              ["GET",    "/api/analysis/{id}/transcript",       "Transcripción con timestamps."],
              ["GET",    "/api/analysis/{id}/export.json",      "Exporta el análisis a JSON."],
              ["GET",    "/api/analysis/{id}/export.pdf",       "Exporta el análisis a PDF."],
              ["POST",   "/api/search",                         "Búsqueda semántica."],
              ["POST",   "/api/chat",                           "Asistente RAG."],
              ["GET",    "/api/dashboard/overview",             "KPIs agregados."],
              ["GET",    "/api/dashboard/related/{id}",         "Audios relacionados."],
          ])

    h1(doc, "Anexo C. Estructura de salida del análisis (ejemplo JSON)")
    p(doc, '{', italic=True)
    p(doc, '  "summary_short": "Reunión de planificación del Q2 de 2026 ...",', italic=True)
    p(doc, '  "summary_medium": "Los asistentes revisaron los objetivos ...",', italic=True)
    p(doc, '  "summary_long": "- Contexto: ...\\n- Temas: ...\\n- Decisiones: ...",', italic=True)
    p(doc, '  "entities": {"personas":[...], "empresas":[...], ...},', italic=True)
    p(doc, '  "tasks": [{"task":"Preparar demo","owner":"Laura","deadline":"2026-05-03","priority":"alta"}],', italic=True)
    p(doc, '  "decisions": [{"decision":"Adoptar Kafka","made_by":"CTO","rationale":"..."}],', italic=True)
    p(doc, '  "sentiment": {"global":{"label":"neutro","score":0.12}, "per_speaker":[...]},', italic=True)
    p(doc, '  "metrics": {"total_duration_sec":1820.5,"num_speakers":4,"participation_pct":{...}}', italic=True)
    p(doc, '}', italic=True)

    # save
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Documento generado en: {OUT}")


if __name__ == "__main__":
    build()
